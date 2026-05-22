from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "corrected-chapters-1-3.docx"


TITLE_LINES = [
    "MIVA OPEN UNIVERSITY",
    "FACULTY OF COMPUTING",
    "DEPARTMENT OF SOFTWARE ENGINEERING",
    "DESIGN AND IMPLEMENTATION OF AN AI-ENHANCED MULTI-TENANT FAMILY FUND MANAGEMENT SYSTEM WITH PREDICTIVE ANALYTICS AND INTELLIGENT REPORTING",
    "BY",
    "AMINU DANLADI HUSSAIN",
    "2024/A/SENG/0156",
    "A PROJECT SUBMITTED TO THE DEPARTMENT OF SOFTWARE ENGINEERING, FACULTY OF COMPUTING, MIVA OPEN UNIVERSITY, IN PARTIAL FULFILMENT OF THE REQUIREMENTS FOR THE AWARD OF THE DEGREE OF BACHELOR OF SCIENCE (B.Sc.) IN SOFTWARE ENGINEERING",
    "SUPERVISOR: DR SAMUEL MAKINDE",
    "MAY, 2026",
]


def set_run_font(run, size: int = 12, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    r_fonts = run._element.get_or_add_rPr().rFonts
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:cs"), "Times New Roman")


def set_paragraph_base(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    paragraph.alignment = alignment
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_after = Pt(8)
    fmt.space_before = Pt(0)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def add_styled_paragraph(document, text: str, *, bold=False, italic=False, size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    paragraph = document.add_paragraph()
    set_paragraph_base(paragraph, alignment)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return paragraph


def add_heading(document, text: str, level: int) -> None:
    style = f"Heading {level}"
    paragraph = document.add_paragraph(style=style)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level in {1, 2} else WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(12 if level >= 3 else 18)
    fmt.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size={1: 16, 2: 14, 3: 12, 4: 12}.get(level, 12), bold=True)


def parse_inline_markdown(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = text.replace("**", "")
    text = text.replace("*", "")
    return text


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [parse_inline_markdown(cell.strip()) for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def add_table(document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.autofit = True
    set_table_borders(table)
    for r_index, row in enumerate(rows):
        for c_index, value in enumerate(row):
            cell = table.cell(r_index, c_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.line_spacing = 1.15
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_run_font(run, size=10 if len(rows[0]) > 3 else 11, bold=r_index == 0)
    document.add_paragraph()


def add_image(document, alt_text: str, image_path: Path) -> None:
    add_styled_paragraph(document, alt_text, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    if image_path.exists():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(6.3))
        document.add_paragraph()
    else:
        add_styled_paragraph(document, f"[Image missing: {image_path.name}]", italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)


def add_markdown_file(document, file_name: str) -> None:
    lines = (ROOT / file_name).read_text().splitlines()
    pending_table: list[str] = []

    def flush_table() -> None:
        nonlocal pending_table
        if pending_table:
            rows = parse_table([line for i, line in enumerate(pending_table) if i != 1])
            add_table(document, rows)
            pending_table = []

    for raw_line in lines:
        line = raw_line.rstrip()
        if pending_table and line.startswith("|"):
            pending_table.append(line)
            continue
        flush_table()

        stripped = line.strip()
        if not stripped or stripped == "---" or stripped.startswith("> **References:**"):
            continue
        if stripped.startswith("# "):
            add_heading(document, stripped[2:].strip(), 1)
            continue
        if stripped.startswith("## "):
            add_heading(document, stripped[3:].strip(), 2)
            continue
        if stripped.startswith("### "):
            add_heading(document, stripped[4:].strip(), 3)
            continue
        if stripped.startswith("#### "):
            add_heading(document, stripped[5:].strip(), 4)
            continue
        if stripped.startswith("|"):
            pending_table.append(line)
            continue
        image_match = re.match(r"!\[(.+?)\]\((.+?)\)", stripped)
        if image_match:
            alt, rel_path = image_match.groups()
            add_image(document, alt, ROOT / rel_path)
            continue
        table_caption = re.match(r"\*Table ([^*]+)\*", stripped)
        if table_caption:
            add_styled_paragraph(document, stripped.strip("*"), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            continue
        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if numbered:
            add_styled_paragraph(document, f"{numbered.group(1)}. {parse_inline_markdown(numbered.group(2))}", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
            continue
        add_styled_paragraph(document, parse_inline_markdown(stripped))
    flush_table()


def add_title_page(document) -> None:
    for index, line in enumerate(TITLE_LINES):
        paragraph = document.add_paragraph()
        set_paragraph_base(paragraph, WD_ALIGN_PARAGRAPH.CENTER)
        paragraph.paragraph_format.space_after = Pt(18 if index in {2, 3, 6, 7, 8} else 10)
        run = paragraph.add_run(line)
        set_run_font(run, size=12, bold=True)
    document.add_page_break()


def configure_styles(document) -> None:
    styles = document.styles
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.color.rgb = RGBColor(0, 0, 0)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    styles["Normal"].font.size = Pt(12)


def add_references(document) -> None:
    document.add_page_break()
    add_heading(document, "REFERENCES", 1)
    references = (ROOT / "references.md").read_text().splitlines()
    for line in references:
        stripped = line.strip()
        if not stripped or stripped.startswith("#") or stripped.startswith(">") or stripped == "---":
            continue
        add_styled_paragraph(document, parse_inline_markdown(stripped), alignment=WD_ALIGN_PARAGRAPH.LEFT)


def main() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    configure_styles(document)
    add_title_page(document)
    for chapter in ["chapter-1.md", "chapter-2.md", "chapter-3.md"]:
        add_markdown_file(document, chapter)
        if chapter != "chapter-3.md":
            document.add_page_break()
    add_references(document)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
